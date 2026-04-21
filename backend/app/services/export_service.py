import os
import uuid
from datetime import datetime
from sqlalchemy.orm import Session
from fastapi import HTTPException

from app.models.translation import Translation
from app.models.export import Export
from app.schemas.export import ExportRequest


# 내보내기 파일 저장 경로
EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


def _get_translation(translation_id: int, db: Session) -> Translation:
    translation = db.query(Translation).filter(Translation.id == translation_id).first()
    if not translation:
        raise HTTPException(status_code=404, detail="해당 번역 기록을 찾을 수 없습니다.")
    return translation


def _generate_filename(fmt: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    uid = uuid.uuid4().hex[:8]
    ext_map = {"pdf": "pdf", "word": "docx", "img": "png"}
    ext = ext_map.get(fmt, fmt)
    return f"zzappago_{timestamp}_{uid}.{ext}"


def export_as_pdf(translation: Translation) -> str:
    """번역 결과를 PDF로 내보내기"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    filename = _generate_filename("pdf")
    filepath = os.path.join(EXPORT_DIR, filename)

    # 한글 폰트 등록 (시스템에 따라 경로 조정 필요)
    font_paths = [
        "C:/Windows/Fonts/malgun.ttf",       # Windows
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",  # Linux
        "/System/Library/Fonts/AppleGothic.ttf",  # macOS
    ]

    font_registered = False
    for font_path in font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("KoreanFont", font_path))
            font_registered = True
            break

    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4

    font_name = "KoreanFont" if font_registered else "Helvetica"

    # 제목
    c.setFont(font_name, 18)
    c.drawString(2 * cm, height - 3 * cm, "ZzapPago Translation")

    # 언어 정보
    c.setFont(font_name, 11)
    c.drawString(2 * cm, height - 4.5 * cm, f"{translation.source_lang} → {translation.target_lang}")

    # 원문
    c.setFont(font_name, 12)
    c.drawString(2 * cm, height - 6.5 * cm, "[원문]")
    c.setFont(font_name, 10)
    y = height - 7.5 * cm
    for line in translation.source_text.split("\n"):
        if y < 3 * cm:
            c.showPage()
            y = height - 3 * cm
            c.setFont(font_name, 10)
        c.drawString(2 * cm, y, line[:80])
        y -= 0.5 * cm

    # 번역문
    y -= 1 * cm
    c.setFont(font_name, 12)
    c.drawString(2 * cm, y, "[번역]")
    y -= 1 * cm
    c.setFont(font_name, 10)
    for line in translation.translated_text.split("\n"):
        if y < 3 * cm:
            c.showPage()
            y = height - 3 * cm
            c.setFont(font_name, 10)
        c.drawString(2 * cm, y, line[:80])
        y -= 0.5 * cm

    c.save()
    return filepath


def export_as_word(translation: Translation) -> str:
    """번역 결과를 Word(docx)로 내보내기"""
    from docx import Document
    from docx.shared import Pt

    filename = _generate_filename("word")
    filepath = os.path.join(EXPORT_DIR, filename)

    doc = Document()

    # 제목
    doc.add_heading("ZzapPago Translation", level=1)

    # 언어 정보
    p = doc.add_paragraph()
    p.add_run(f"{translation.source_lang} → {translation.target_lang}").font.size = Pt(11)

    doc.add_paragraph("")

    # 원문
    doc.add_heading("원문", level=2)
    doc.add_paragraph(translation.source_text)

    # 번역문
    doc.add_heading("번역", level=2)
    doc.add_paragraph(translation.translated_text)

    # 날짜
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(f"생성일: {translation.created_at.strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)

    doc.save(filepath)
    return filepath


def export_as_img(translation: Translation) -> str:
    """번역 결과를 이미지(PNG)로 내보내기"""
    from PIL import Image, ImageDraw, ImageFont

    filename = _generate_filename("img")
    filepath = os.path.join(EXPORT_DIR, filename)

    width, height = 800, 600

    img = Image.new("RGB", (width, height), color="#FFFFFF")
    draw = ImageDraw.Draw(img)

    # 한글 폰트 로드
    font_paths = [
        "C:/Windows/Fonts/malgun.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/System/Library/Fonts/AppleGothic.ttf",
    ]

    title_font = None
    body_font = None
    small_font = None

    for font_path in font_paths:
        if os.path.exists(font_path):
            title_font = ImageFont.truetype(font_path, 28)
            body_font = ImageFont.truetype(font_path, 16)
            small_font = ImageFont.truetype(font_path, 12)
            break

    if not title_font:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    # 배경 헤더
    draw.rectangle([0, 0, width, 70], fill="#1ec800")
    draw.text((30, 18), "ZzapPago Translation", fill="#FFFFFF", font=title_font)

    # 언어 정보
    draw.text((30, 90), f"{translation.source_lang} → {translation.target_lang}", fill="#888888", font=small_font)

    # 원문
    draw.text((30, 120), "[ 원문 ]", fill="#333333", font=body_font)
    source_lines = translation.source_text[:300].split("\n")
    y = 150
    for line in source_lines[:6]:
        draw.text((30, y), line[:60], fill="#333333", font=body_font)
        y += 24

    # 구분선
    y += 10
    draw.line([(30, y), (width - 30, y)], fill="#E5E7EB", width=1)
    y += 20

    # 번역문
    draw.text((30, y), "[ 번역 ]", fill="#1ec800", font=body_font)
    y += 30
    translated_lines = translation.translated_text[:300].split("\n")
    for line in translated_lines[:6]:
        draw.text((30, y), line[:60], fill="#333333", font=body_font)
        y += 24

    img.save(filepath, "PNG")
    return filepath


def create_export(req: ExportRequest, db: Session) -> dict:
    """내보내기 실행 및 DB 기록"""
    translation = _get_translation(req.translation_id, db)

    if req.format == "pdf":
        filepath = export_as_pdf(translation)
    elif req.format == "word":
        filepath = export_as_word(translation)
    elif req.format == "img":
        filepath = export_as_img(translation)
    else:
        raise HTTPException(status_code=400, detail=f"지원하지 않는 형식: {req.format}")

    export_record = Export(
        translation_id=req.translation_id,
        format=req.format,
        file_path=filepath,
    )
    db.add(export_record)
    db.commit()
    db.refresh(export_record)

    return {
        "id": export_record.id,
        "translation_id": export_record.translation_id,
        "format": export_record.format,
        "file_path": export_record.file_path,
        "created_at": export_record.created_at,
    }
